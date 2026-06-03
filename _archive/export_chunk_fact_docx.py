from __future__ import annotations

import argparse
import json
from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


def _xml_safe(text: Any) -> str:
    value = str(text or "")
    return "".join(
        char
        for char in value
        if char in "\t\n\r" or ord(char) >= 0x20
    )


def _load_json(path: str | Path) -> Any:
    with open(path, "r", encoding="utf-8") as handle:
        return json.load(handle)


def _facts_from_payload(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, dict):
        facts = payload.get("facts", [])
        return facts if isinstance(facts, list) else []
    return payload if isinstance(payload, list) else []


def _raw_name(fact: dict[str, Any]) -> str:
    raw = fact.get("raw") if isinstance(fact.get("raw"), dict) else {}
    return str(raw.get("raw_name") or fact.get("metric") or "")


def _metric_core(fact: dict[str, Any]) -> str:
    raw = fact.get("raw") if isinstance(fact.get("raw"), dict) else {}
    return str(raw.get("metric_core") or fact.get("metric") or "")


def _fact_decision(fact: dict[str, Any]) -> str:
    return str(fact.get("decision") or "")


def _keep_fact(fact: dict[str, Any]) -> bool:
    return _fact_decision(fact).lower() in {"keep", "rescue"}


def _short(text: Any, max_chars: int = 900) -> str:
    clean = " ".join(_xml_safe(text).split())
    if len(clean) <= max_chars:
        return clean
    return clean[: max_chars - 1].rstrip() + "..."


def _add_kv(paragraph, label: str, value: Any) -> None:
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(value or ""))


def _set_doc_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Aptos"


def export_docx(
    *,
    chunks_path: str | Path,
    pass1_path: str | Path,
    pass2_path: str | Path,
    output_path: str | Path,
    title: str,
) -> dict[str, int]:
    chunks = _load_json(chunks_path)
    pass1_facts = _facts_from_payload(_load_json(pass1_path))
    pass2_facts = _facts_from_payload(_load_json(pass2_path))
    pass2_by_fact_id = {str(fact.get("fact_id") or ""): fact for fact in pass2_facts}

    kept_facts_by_chunk: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for fact in pass1_facts:
        if _keep_fact(fact):
            kept_facts_by_chunk[str(fact.get("chunk_id") or "")].append(fact)

    kept_chunks = [
        chunk
        for chunk in chunks
        if kept_facts_by_chunk.get(str(chunk.get("chunk_id") or ""))
    ]

    document = Document()
    _set_doc_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    document.add_heading(title, level=1)
    summary = document.add_paragraph()
    _add_kv(summary, "Chunks with kept facts", len(kept_chunks))
    summary.add_run(" | ")
    _add_kv(summary, "Kept/rescued facts", sum(len(v) for v in kept_facts_by_chunk.values()))
    summary.add_run(" | ")
    _add_kv(summary, "Source chunks", Path(chunks_path).name)

    for index, chunk in enumerate(kept_chunks, start=1):
        if index > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        chunk_id = str(chunk.get("chunk_id") or "")
        page_start = chunk.get("page_start")
        page_end = chunk.get("page_end")
        document.add_heading(f"{index}. {chunk_id}", level=2)

        meta = document.add_paragraph()
        _add_kv(meta, "Pages", f"{page_start}-{page_end}" if page_start != page_end else page_start)
        meta.add_run(" | ")
        _add_kv(meta, "Section", chunk.get("section_title"))
        meta.add_run(" | ")
        _add_kv(meta, "Facts", len(kept_facts_by_chunk[chunk_id]))

        document.add_heading("Kept Paragraph / Chunk Text", level=3)
        document.add_paragraph(_short(chunk.get("content"), 3500))

        document.add_heading("Extracted Facts and Stored Normalization", level=3)
        table = document.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = [
            "Raw fact",
            "Core",
            "Value",
            "Period",
            "Evidence",
            "Stored decision",
            "Canonical / proposed",
            "Score / note",
        ]
        for cell, header in zip(table.rows[0].cells, headers):
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for fact in kept_facts_by_chunk[chunk_id]:
            stored = pass2_by_fact_id.get(str(fact.get("fact_id") or ""), {})
            raw = fact.get("raw") if isinstance(fact.get("raw"), dict) else {}
            row = table.add_row().cells
            row[0].text = _short(_raw_name(fact), 160)
            row[1].text = _short(_metric_core(fact), 160)
            row[2].text = _short(f"{fact.get('value') or raw.get('raw_value') or ''} {fact.get('unit') or raw.get('raw_unit') or ''}".strip(), 120)
            row[3].text = _short(fact.get("period") or raw.get("raw_period") or "", 120)
            row[4].text = _short(fact.get("evidence") or raw.get("source_sentence") or "", 300)
            row[5].text = _short(
                stored.get("normalization_decision")
                or stored.get("decision")
                or "",
                120,
            )
            row[6].text = _short(
                stored.get("canonical_id")
                or stored.get("proposed_canonical_id")
                or "",
                160,
            )
            row[7].text = _short(
                " | ".join(
                    part
                    for part in (
                        f"confidence={stored.get('mapping_confidence')}" if stored.get("mapping_confidence") else "",
                        f"note={stored.get('mapping_note')}" if stored.get("mapping_note") else "",
                    )
                    if part
                ),
                240,
            )

    output_path = Path(output_path)
    document.save(output_path)
    return {
        "source_chunks": len(chunks),
        "chunks_with_kept_facts": len(kept_chunks),
        "kept_or_rescued_facts": sum(len(v) for v in kept_facts_by_chunk.values()),
        "pass2_facts": len(pass2_facts),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="Export chunk text, extracted facts, and stored normalization to DOCX.")
    parser.add_argument("--chunks", required=True)
    parser.add_argument("--pass1", required=True)
    parser.add_argument("--pass2", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--title", default="Chunk Fact Audit")
    args = parser.parse_args()
    result = export_docx(
        chunks_path=args.chunks,
        pass1_path=args.pass1,
        pass2_path=args.pass2,
        output_path=args.output,
        title=args.title,
    )
    print(json.dumps(result, indent=2))
    print(f"Wrote {Path(args.output).resolve()}")


if __name__ == "__main__":
    main()
